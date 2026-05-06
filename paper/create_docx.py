#!/usr/bin/env python3
"""Generate journal-style DOCX from paper content with embedded figures."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
FIGURES_DIR = str(PROJECT_DIR / 'data' / 'processed' / 'figures')

def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_figure(doc, image_path, caption, fig_num, width=5.5):
    """Add a figure with caption."""
    if not os.path.exists(image_path):
        p = doc.add_paragraph(f"[Figure {fig_num}: {caption} - Image not found: {image_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f'Figure {fig_num}. ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run = cap.add_run(caption)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    doc.add_paragraph()

def add_table(doc, caption, headers, data, bold_last_row=False):
    """Add a table with caption."""
    table_caption = doc.add_paragraph()
    cap_parts = caption.split('.', 1)
    run = table_caption.add_run(f'{cap_parts[0]}. ')
    run.bold = True
    run.font.name = 'Times New Roman'
    if len(cap_parts) > 1:
        run = table_caption.add_run(cap_parts[1].strip())
        run.font.name = 'Times New Roman'
    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(cell, 'D9D9D9')

    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if bold_last_row and row_idx == len(data) - 1:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    return table

def add_paragraph(doc, text, indent=True):
    """Add a justified paragraph."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p

def add_section_heading(doc, text, level=1):
    """Add a section heading."""
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(12)
    else:
        run.italic = True
    run.font.name = 'Times New Roman'
    return h

def create_paper():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Graph Neural Networks for Seizure Onset Zone Localization in Drug-Resistant Epilepsy: A Multi-Site Intracranial EEG Study')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    # Running title
    running = doc.add_paragraph()
    running.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = running.add_run('Running title: GNN for SOZ Localization')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph()

    # Authors with superscript
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run('Abtin Akhtari')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = authors.add_run('1,*')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.superscript = True

    # Affiliation
    affil = doc.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = affil.add_run('1')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.superscript = True
    run = affil.add_run(' University of Nebraska Medical Center, Omaha, NE, USA')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

    # Corresponding author
    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = corr.add_run('* ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run = corr.add_run('Corresponding author: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    run = corr.add_run('Abtin Akhtari (abtin.akhtary@gmail.com)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    # Word count
    wc = doc.add_paragraph()
    wc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = wc.add_run('Word count: ~4,500 | Figures: 10 | Tables: 8')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)

    doc.add_paragraph()

    # Abstract
    abs_head = doc.add_paragraph()
    abs_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = abs_head.add_run('Abstract')
    run.bold = True
    run.font.name = 'Times New Roman'

    abstract_text = """Finding the seizure onset zone (SOZ) accurately can make or break epilepsy surgery outcomes. Intracranial EEG remains our best tool for this job, but having clinicians manually review these recordings takes considerable time and different experts often disagree on their interpretations. We tackled this problem by building a graph neural network (GNN) that treats electrode arrays as interconnected networks—each electrode becomes a node, and the functional connections between them form the edges. We pooled data from two publicly available datasets out of the Hospital of the University of Pennsylvania: 81 patients total, 212 recordings, with 126 having confirmed SOZ labels from clinical experts. Our approach uses GraphSAGE with self-supervised pretraining and data augmentation during training. The base model reached an AUC of 0.768, which jumped to 0.785 ± 0.013 once we added Random Forest predictions as an extra feature for each node. What stood out most was the sensitivity: our model caught 63% of true SOZ electrodes, compared to just 44% for Random Forest alone—even though RF scored higher on overall AUC (0.847). This matters clinically because missing an SOZ electrode during surgery planning is far worse than flagging a few extra ones for review. These results point toward GNN-based tools, especially hybrid versions that blend classical ML with graph learning, as promising aids for surgical planning in epilepsy."""

    abs_para = doc.add_paragraph(abstract_text)
    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_para.paragraph_format.first_line_indent = Inches(0.5)
    for run in abs_para.runs:
        run.font.name = 'Times New Roman'

    # Keywords
    kw = doc.add_paragraph()
    run = kw.add_run('Keywords: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run = kw.add_run('seizure onset zone, graph neural networks, intracranial EEG, epilepsy surgery, deep learning')
    run.font.name = 'Times New Roman'
    run.italic = True

    doc.add_paragraph()

    # Section 1: Introduction
    add_section_heading(doc, '1. Introduction')

    intro_paras = [
        "Around 50 million people live with epilepsy globally. For roughly a third of them, medications simply don't control the seizures (Kwan & Brodie, 2000). Surgery offers these patients their best shot at becoming seizure-free—success rates hover between 50% and 80%, depending on what's causing the epilepsy (Engel et al., 2012). But here's the catch: surgery only works if we can pinpoint exactly where seizures start. Get the SOZ wrong, and the operation fails.",
        "Intracranial EEG, whether through depth electrodes (SEEG) or surface grids (ECoG), gives us recordings straight from brain tissue with excellent spatial and temporal detail. Epileptologists spend hours poring over these signals, looking for the telltale patterns of seizure onset. It's painstaking work, and frankly, subjective. Ask three experts to mark the SOZ and you might get three different answers (Varatharajah et al., 2018). We clearly need more objective, automated approaches.",
        "Deep learning has made inroads into neurophysiological signal analysis over the past few years. Standard convolutional and recurrent networks, though, weren't really designed for electrode arrays. These arrays have irregular spatial layouts that change from patient to patient based on where clinicians think the seizures might originate. Graph neural networks (GNNs) handle this naturally—they represent electrodes as nodes and encode the functional relationships between recording sites as edges (Bessadok et al., 2022). It's a framework that fits the problem well.",
        "In this paper, we describe a GNN pipeline for SOZ localization that we built and tested on 81 patients from two open-access iEEG repositories. We make five main contributions: first, a preprocessing pipeline that works with multiple data formats; second, a self-supervised pretraining scheme that makes use of unlabeled recordings; third, a thorough comparison of different augmentation strategies; fourth, head-to-head testing of several GNN architectures; and fifth, an exploration of hybrid methods that pair traditional machine learning with graph-based deep learning."
    ]

    for text in intro_paras:
        add_paragraph(doc, text)

    # Section 2: Materials and Methods
    add_section_heading(doc, '2. Materials and Methods')
    add_section_heading(doc, '2.1 Datasets', level=2)

    add_paragraph(doc, "We worked with two datasets from the Hospital of the University of Pennsylvania (HUP), both freely available on OpenNeuro:")
    add_paragraph(doc, "Dataset 1 (ds003029): Contains recordings from 23 drug-resistant epilepsy patients who went through presurgical workups. The files come in BrainVision format, and the SOZ annotations are tucked into the BIDS events.tsv files as event markers.")
    add_paragraph(doc, "Dataset 2 (ds004100): A larger collection with 58 patients who had either SEEG or ECoG monitoring. These recordings use European Data Format (EDF), and the SOZ labels appear directly in the channels.tsv metadata. All these patients eventually had surgery or laser ablation, with Engel outcome scores on record.")
    add_paragraph(doc, "Altogether, we had 212 recordings to work with. Of these, 126 came with expert-annotated SOZ labels (Table 1). The rest we used for pretraining.")

    add_table(doc, 'Table 1. Dataset characteristics',
              ['Characteristic', 'ds003029', 'ds004100', 'Combined'],
              [
                  ['Patients', '23', '58', '81'],
                  ['Recordings', '103', '109', '212'],
                  ['Recordings with SOZ', '30', '96', '126'],
                  ['Electrode type', 'SEEG/ECoG', 'SEEG/ECoG', 'Mixed'],
                  ['Sampling rate', '500-2000 Hz', '500 Hz', 'Variable']
              ])

    # 2.2 Preprocessing
    add_section_heading(doc, '2.2 Preprocessing Pipeline', level=2)

    add_paragraph(doc, "We built our preprocessing to handle both BrainVision and EDF files. Four steps:")
    add_paragraph(doc, "Extracting epochs: We found the seizure onset timestamp in each recording and pulled out 60 seconds of data centered on that moment—30 seconds before onset, 30 seconds after. Any channels flagged as bad in the metadata got tossed.")
    add_paragraph(doc, "Notch filtering: Power line noise at 60 Hz and its harmonics (120, 180, 240 Hz) had to go. We included the fourth harmonic because 240 Hz sits close to our bandpass cutoff and could sneak through otherwise.")
    add_paragraph(doc, "Common average referencing: Volume conduction and common-mode noise muddy up iEEG signals. To clean this up, we subtracted the mean voltage across all electrodes at each time point.")
    add_paragraph(doc, "Bandpass filtering: We applied an elliptic filter with corners at 1 Hz and 250 Hz. This keeps both the slow rhythms and the high-frequency oscillations (HFOs) that researchers have linked to epileptogenic tissue (Jacobs et al., 2012).")

    add_figure(doc, f'{FIGURES_DIR}/preprocessing_demo.png',
               'Preprocessing steps: (A) raw signal, (B) after notch filtering, (C) after common average reference, (D) after bandpass filtering.',
               1, width=5.5)

    # 2.3 Feature Extraction
    add_section_heading(doc, '2.3 Feature Extraction', level=2)

    add_paragraph(doc, "After preprocessing, we chopped each recording into 12-second windows with no overlap. From each window, we pulled out features for every electrode:")
    add_paragraph(doc, "Spectral features (6 total): Power in the delta (1-4 Hz), theta (4-8 Hz), alpha (8-12 Hz), beta (15-25 Hz), low gamma (35-50 Hz), and HFO (80-150 Hz) bands. We used Welch's method for the power estimates.")
    add_paragraph(doc, "Statistical features (4 total): Variance, line length (summing up the absolute differences between consecutive samples), kurtosis, and skewness.")
    add_paragraph(doc, "Relative HFO power (1): The ratio of HFO power to total power. Elevated HFO activity tends to show up in epileptogenic regions.")
    add_paragraph(doc, "We then aggregated across windows using the mean, standard deviation, and maximum—giving us 33 features per electrode. To capture functional connectivity, we computed Pearson correlations between electrode pairs within each window, then aggregated those the same way.")

    # 2.4 Graph Construction
    add_section_heading(doc, '2.4 Graph Construction', level=2)

    add_paragraph(doc, "Each recording became a graph. Electrodes are nodes, and we drew edges between any pair whose correlation topped 0.32 (a threshold we landed on through hyperparameter tuning). The 33 aggregated features went into each node, while edge weights encoded connectivity strength.")
    add_paragraph(doc, "For labels, any electrode that clinicians had marked as SOZ got a positive label; the rest were negative. The class balance was pretty skewed—only about 8.9% of electrodes fell in the SOZ category across our combined dataset.")

    add_figure(doc, f'{FIGURES_DIR}/graph_demo.png',
               'Example graph from an iEEG array. Nodes (electrodes) colored by SOZ label; edges show functional connectivity above threshold.',
               2, width=4.5)

    # 2.5 Model Architecture
    add_section_heading(doc, '2.5 Model Architecture', level=2)

    add_paragraph(doc, "We went with GraphSAGE (Hamilton et al., 2017), which builds node representations by pulling in information from neighboring nodes. Our encoder stacks two GraphSAGE layers with batch normalization, 128 hidden dimensions, and 21% dropout.")
    add_paragraph(doc, "On top of that sits a classification head: a small MLP with 64 hidden units, 55% dropout, and a sigmoid output for the binary SOZ prediction.")

    # 2.6 Training Strategy
    add_section_heading(doc, '2.6 Training Strategy', level=2)

    add_paragraph(doc, "Pretraining without labels: Since we had plenty of unlabeled recordings, we pretrained the encoder on a forecasting task—predicting the features at time window t+1 from window t. This ran on 827 window pairs drawn from all available data.")
    add_paragraph(doc, "Fine-tuning with labels: We then trained the full model for SOZ classification using weighted binary cross-entropy. The weights were set inversely to class frequencies to help with the imbalance problem.")
    add_paragraph(doc, "Augmentation on the fly: During training, we randomly dropped 7.6% of edges, added Gaussian noise (σ = 0.011) to features, and masked out 10% of feature values.")
    add_paragraph(doc, "Optimization details: AdamW optimizer, learning rates of 1.2×10⁻⁴ for pretraining and 6.1×10⁻⁴ for fine-tuning, cosine annealing schedule, and early stopping based on validation AUC.")

    # 2.7 Experimental Setup
    add_section_heading(doc, '2.7 Experimental Setup', level=2)

    add_paragraph(doc, "We split the data by patient—no patient appeared in more than one split—with 60% for training, 20% for validation, and 20% for testing. That worked out to 70 training graphs, 36 validation graphs, and 20 test graphs.")
    add_paragraph(doc, "We tried three GNN flavors: GraphSAGE, Graph Attention Network (GAT; Velickovic et al., 2018), and Graph Isomorphism Network (GIN; Xu et al., 2019). Optuna handled the hyperparameter search, running 50 trials for each architecture.")

    add_figure(doc, f'{FIGURES_DIR}/optuna_optimization.png',
               'Optuna optimization history across 50 Bayesian search trials.',
               3, width=5.0)

    # Section 3: Results
    add_section_heading(doc, '3. Results')

    # 3.1 Baseline Comparison
    add_section_heading(doc, '3.1 Comparison with Classical Machine Learning', level=2)

    add_paragraph(doc, "To put our GNN results in context, we trained several standard ML models on the same node features, just without the graph structure. Table 2 shows what happened.")

    add_table(doc, 'Table 2. Comparison with classical ML baselines (combined dataset)',
              ['Method', 'Test AUC', 'Test F1', 'Test Precision', 'Test Recall'],
              [
                  ['Random Forest', '0.847', '0.373', '0.326', '0.436'],
                  ['Logistic Regression', '0.741', '0.336', '0.238', '0.571'],
                  ['Lasso (L1)', '0.738', '0.338', '0.240', '0.571'],
                  ['Ridge (L2)', '0.741', '0.336', '0.238', '0.571'],
                  ['LR (tuned, CV)', '0.742', '0.340', '0.242', '0.571'],
                  ['SVM (RBF kernel)', '0.713', '0.097', '0.200', '0.064'],
                  ['GraphSAGE (ours)', '0.763 ± 0.008', '0.273 ± 0.007', '0.175 ± 0.006', '0.623 ± 0.014']
              ], bold_last_row=True)

    add_paragraph(doc, "Here's what surprised us: Random Forest actually beat the GNN on AUC, hitting 0.847 versus our 0.763. But look at the recall numbers. The GNN caught 62.3% of SOZ electrodes while Random Forest only managed 43.6%—that's a 43% relative improvement. When the goal is to avoid missing any true SOZ contacts, that difference matters enormously.")

    # 3.2 GNN Architecture Comparison
    add_section_heading(doc, '3.2 GNN Architecture Comparison', level=2)

    add_paragraph(doc, "We benchmarked different GNN architectures on just the ds003029 dataset first (Table 3). After tuning, GraphSAGE came out on top with a test AUC of 0.730, beating GAT (0.702) and GIN (0.562) by a fair margin.")

    add_table(doc, 'Table 3. Performance comparison of GNN architectures (ds003029 only)',
              ['Architecture', 'Val AUC', 'Test AUC', 'Test F1'],
              [
                  ['GraphSAGE (baseline)', '0.920', '0.697', '0.478'],
                  ['GraphSAGE (tuned)', '0.983', '0.730', '0.404'],
                  ['GAT (tuned)', '0.925', '0.702', '0.411'],
                  ['GIN', '0.660', '0.562', '0.000']
              ])

    # 3.3 Augmentation
    add_section_heading(doc, '3.3 Data Augmentation Analysis', level=2)

    add_paragraph(doc, "Not all augmentation helps. Table 4 breaks down what we found. Online augmentation—dropping edges, adding noise, masking features—pushed the test AUC from 0.730 up to 0.761, a solid 4.2% gain. Time-shift augmentation, mixup, and SMOTE either did nothing or actively hurt performance.")

    add_table(doc, 'Table 4. Effect of data augmentation techniques',
              ['Technique', 'Test AUC', 'Change'],
              [
                  ['None (baseline)', '0.730', '—'],
                  ['Online augmentation', '0.761', '+4.2%'],
                  ['Time-shift', '0.718', '-1.6%'],
                  ['Mixup', '0.731', '+0.1%'],
                  ['SMOTE', '0.696', '-4.7%']
              ])

    # 3.4 Combined Dataset
    add_section_heading(doc, '3.4 Combined Dataset Performance', level=2)

    add_paragraph(doc, "Pooling both datasets gave us more training data—70 labeled graphs instead of roughly 20—and the numbers improved. A single training run hit 0.768 AUC with 65% recall. Across five random seeds, we averaged 0.763 ± 0.008 AUC and 62.3% ± 1.4% recall (Table 5).")

    add_table(doc, 'Table 5. Performance on combined dataset',
              ['Metric', 'Single Run', '5 Seeds (mean ± std)'],
              [
                  ['Validation AUC', '0.751', '—'],
                  ['Test AUC', '0.768', '0.763 ± 0.008'],
                  ['Test F1', '0.296', '0.273 ± 0.007'],
                  ['Test Precision', '0.192', '0.175 ± 0.006'],
                  ['Test Recall', '0.650', '0.623 ± 0.014']
              ])

    add_table(doc, 'Table 5a. GNN per-seed results on combined dataset',
              ['Seed', 'Test AUC', 'Test F1', 'Test Recall'],
              [
                  ['42', '0.753', '0.260', '0.621'],
                  ['123', '0.776', '0.278', '0.643'],
                  ['456', '0.762', '0.281', '0.629'],
                  ['789', '0.756', '0.270', '0.600'],
                  ['2024', '0.768', '0.274', '0.621']
              ])

    add_figure(doc, f'{FIGURES_DIR}/roc_curve.png',
               'ROC curve showing discrimination performance on the test set. GraphSAGE reached 0.768 AUC on the combined dataset.',
               4, width=4.5)

    # 3.5 Summary
    add_section_heading(doc, '3.5 Comparison with Single-Dataset Training', level=2)

    add_paragraph(doc, "Table 6 traces our progress through the experiments. More data helped. Augmentation helped. But the biggest jump came from the hybrid approach.")

    add_table(doc, 'Table 6. Summary of experimental progression',
              ['Configuration', 'Training Graphs', 'Test AUC', 'Test Recall'],
              [
                  ['ds003029 baseline', '~20', '0.697', '0.47'],
                  ['ds003029 tuned', '~20', '0.730', '0.40'],
                  ['ds003029 + augmentation', '~20', '0.761', '0.52'],
                  ['Combined dataset (single)', '70', '0.768', '0.650'],
                  ['Combined dataset (5 seeds)', '70', '0.763 ± 0.008', '0.623 ± 0.014'],
                  ['Combined + RF feature', '70', '0.785 ± 0.013', '0.630 ± 0.052']
              ], bold_last_row=True)

    # 3.6 Hybrid Approaches
    add_section_heading(doc, '3.6 Hybrid Approaches to Improve AUC', level=2)

    add_paragraph(doc, "Random Forest's AUC lead (0.847 vs. 0.763) bugged us, so we tried a few ways to close the gap without giving up the GNN's recall advantage.")
    add_paragraph(doc, "Does graph structure actually help? We trained an MLP with the same architecture but no message passing—just treating each electrode independently. It scored 0.742 ± 0.009 AUC, compared to 0.763 ± 0.008 for the GNN. So yes, the graph adds about 2.8% on top.")
    add_paragraph(doc, "What about different graph sparsity? We tried correlation thresholds of 0.4, 0.5, and 0.6 to make sparser graphs. All of them performed worse (AUC between 0.72 and 0.75). The denser 0.32 threshold seems to work better.")
    add_paragraph(doc, "Stacking the models? Combining RF and GNN predictions through a meta-learner got us to 0.837 AUC, almost matching RF alone. But recall dropped to 40%. The ensemble basically learned to trust RF's conservative predictions.")
    add_paragraph(doc, "RF predictions as a feature? This worked best. We ran Random Forest first, then fed its probability outputs as an extra node feature into the GNN. The result: 0.785 ± 0.013 AUC with 63% ± 5% recall. We got most of the AUC boost while keeping the GNN's sensitivity edge (Table 7).")

    add_table(doc, 'Table 7. Comparison of hybrid approaches',
              ['Method', 'Test AUC', 'Test F1', 'Test Precision', 'Test Recall'],
              [
                  ['Random Forest', '0.847', '0.373', '0.326', '0.436'],
                  ['GNN (baseline)', '0.763 ± 0.008', '0.273 ± 0.007', '0.175 ± 0.006', '0.623 ± 0.014'],
                  ['MLP (no graph)', '0.742 ± 0.009', '0.286 ± 0.007', '0.192 ± 0.007', '0.566 ± 0.015'],
                  ['Stacking (RF + GNN)', '0.837', '0.336', '0.290', '0.400'],
                  ['Stacking (full)', '0.829', '0.318', '0.292', '0.350'],
                  ['GNN + RF feature', '0.785 ± 0.013', '0.293 ± 0.012', '0.192 ± 0.015', '0.630 ± 0.052']
              ], bold_last_row=True)

    add_paragraph(doc, "The RF-feature hybrid strikes a nice balance. It taps into Random Forest's strong per-node discrimination while letting the GNN refine things based on how electrodes connect to each other.")

    add_figure(doc, f'{FIGURES_DIR}/confusion_matrix.png',
               'Test set confusion matrix showing true/false positive and negative counts.',
               5, width=4.0)

    # Section 4: Discussion
    add_section_heading(doc, '4. Discussion')

    add_paragraph(doc, "Our GNN pipeline can identify SOZ electrodes from iEEG with clinically useful accuracy. The 0.768 AUC we achieved sits comfortably within the 0.65–0.85 range that prior ML studies have reported for this task (Varatharajah et al., 2018; Bernabei et al., 2022).")

    add_section_heading(doc, '4.1 Clinical Implications', level=2)

    add_paragraph(doc, "What jumps out is the 62.3% recall. In surgical planning, you really don't want to miss SOZ electrodes—incomplete resection means the seizures come back. Random Forest beat us on overall discrimination (0.847 vs. 0.763 AUC), but it only caught 43.6% of the true SOZ contacts. That 43% relative improvement in sensitivity could translate to better surgical outcomes.")
    add_paragraph(doc, "Why does the GNN do better at finding true positives? We think it comes down to connectivity. SOZ electrodes tend to have distinctive patterns in how they link to their neighbors—patterns that get lost when you analyze each electrode in isolation, the way Random Forest does.")
    add_paragraph(doc, "The flip side is low precision (17.5%). Our model throws up a lot of false alarms. But in practice, that's manageable. Clinicians can review flagged electrodes and discard the false positives. What they can't do is go back and find electrodes the model missed entirely.")

    add_section_heading(doc, '4.2 Understanding the Performance Gap', level=2)

    add_paragraph(doc, "We honestly didn't expect Random Forest to beat the GNN on AUC. Graphs are supposed to help, right? But after digging in, we think a few factors explain it.")
    add_paragraph(doc, "First, sample size. We trained on just 70 graphs with about 6,200 labeled nodes total. That's not much by deep learning standards. Neural networks tend to memorize training patterns when data is tight (Shorten & Khoshgoftaar, 2019). Random Forest handles small datasets more gracefully—its bagging and feature subsampling naturally resist overfitting (Katzmann et al., 2020).")
    add_paragraph(doc, "Second, our features are already pretty good. Band powers, HFO activity, line length, statistical moments—these biomarkers encode decades of clinical knowledge. When the input features carry that much signal, a simple model can exploit them directly. The extra complexity of message passing may not buy you much (Chen et al., 2020).")
    add_paragraph(doc, "Third, graph construction isn't a solved problem. We defined edges based on correlation thresholds, but who's to say that captures the connectivity patterns that matter? If the graph structure is noisy, the GNN might just be propagating garbage through its neighborhood aggregation.")
    add_paragraph(doc, "Still, the GNN's recall advantage is real and consistent. Even if RF discriminates better overall, it does so by being conservative. The GNN picks up on patterns that help identify true SOZ contacts—at the cost of more false positives. For a screening tool, that tradeoff makes sense.")

    add_section_heading(doc, '4.3 Importance of Multi-Site Data', level=2)

    add_paragraph(doc, "Pooling two datasets, even from the same institution with different protocols, helped more than any augmentation trick we tried. This underscores how much the epilepsy imaging community could gain from sharing data. Platforms like OpenNeuro make that possible.")

    add_section_heading(doc, '4.4 Limitations and Future Directions', level=2)

    add_paragraph(doc, "A few caveats. All our data came from one institution—we don't yet know how well this generalizes elsewhere. The dataset is still modest by deep learning standards. And we evaluated at the electrode level, not the patient level, which might paint an overly rosy picture.")
    add_paragraph(doc, "Going forward, we'd like to test on data from other centers, bring in anatomical information, and extend to seizure prediction. Explainability tools for GNNs could also reveal which features and connectivity patterns drive the model's decisions.")

    # Section 5: Conclusion
    add_section_heading(doc, '5. Conclusion')

    add_paragraph(doc, "We built a GNN framework for automated SOZ localization in patients undergoing intracranial EEG monitoring for drug-resistant epilepsy. Using two public datasets (81 patients, 126 labeled recordings), self-supervised pretraining, and online augmentation, our GraphSAGE model achieved 0.768 AUC with 62–65% recall. Random Forest scored higher on AUC (0.847), but the GNN caught far more true SOZ electrodes (62.3% vs. 43.6%)—a critical advantage when surgical success depends on not missing any.")
    add_paragraph(doc, "Our experiments with hybrid methods turned up three key findings. First, graph structure genuinely helps—there's a 2.8% AUC gain over a comparable MLP. Second, stacking RF and GNN predictions boosts AUC but tanks recall. Third, feeding RF's predictions as an extra node feature gives us both: 0.785 ± 0.013 AUC with 63% ± 5% recall.")
    add_paragraph(doc, "The takeaway? Graph-based deep learning, especially when combined with classical ML, shows real promise as a decision-support tool for epilepsy surgery. The next steps are prospective validation, testing across institutions, and figuring out how to fit these models into clinical workflows.")

    # Data Availability
    add_section_heading(doc, 'Data and Code Availability')

    add_paragraph(doc, "Both datasets are publicly available on OpenNeuro: ds003029 (https://openneuro.org/datasets/ds003029) and ds004100 (https://openneuro.org/datasets/ds004100). Our code is at: https://github.com/abtinunmc/GNN-PROJECT")

    # References
    doc.add_page_break()
    ref_head = doc.add_paragraph()
    run = ref_head.add_run('References')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    references = [
        "Bernabei, J. M., et al. (2022). Normative intracranial EEG maps epileptogenic tissues in focal epilepsy. Brain, 145(6), 1949-1961.",
        "Bessadok, A., Mahjoub, M. A., & Rekik, I. (2022). Graph neural networks in network neuroscience. IEEE TPAMI, 44(7), 3847-3867.",
        "Chen, T., et al. (2020). Big self-supervised models are strong semi-supervised learners. NeurIPS, 33, 22243-22255.",
        "Engel, J., et al. (2012). Early surgical therapy for drug-resistant temporal lobe epilepsy. JAMA, 307(9), 922-930.",
        "Hamilton, W. L., Ying, R., & Leskovec, J. (2017). Inductive representation learning on large graphs. NeurIPS, 30, 1024-1034.",
        "Jacobs, J., et al. (2012). High-frequency oscillations in clinical epilepsy. Progress in Neurobiology, 98(3), 302-315.",
        "Katzmann, A., et al. (2020). Deep random forests for small sample size prediction. IEEE ISBI, 1543-1547.",
        "Kwan, P., & Brodie, M. J. (2000). Early identification of refractory epilepsy. NEJM, 342(5), 314-319.",
        "Shorten, C., & Khoshgoftaar, T. M. (2019). A survey on image data augmentation for deep learning. J Big Data, 6(1), 60.",
        "Varatharajah, Y., et al. (2018). Integrating AI with real-time iEEG monitoring. J Neural Eng, 15(4), 046035.",
        "Velickovic, P., et al. (2018). Graph attention networks. ICLR.",
        "Whalen, S., et al. (2022). Navigating the pitfalls of applying ML in genomics. Nature Rev Genetics, 23(3), 169-181.",
        "Xu, K., et al. (2019). How powerful are graph neural networks? ICLR."
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    # Save
    output_path = SCRIPT_DIR / 'SOZ_Localization_GNN_Paper_v2.docx'
    doc.save(str(output_path))
    print(f"Paper saved to: {output_path}")

if __name__ == '__main__':
    create_paper()
